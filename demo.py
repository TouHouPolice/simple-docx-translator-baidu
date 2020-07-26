
 

import hashlib
import urllib
import random
import requests
import easygui as eg
import docx2txt
import docx
import time
import sys
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText#'\n'.join(fullText)

appid = '20200307000394085' #你的appid
secretKey = 'qioMonQVhuyat02fNvmw' #你的密钥

targetFile=""
targetFile=eg.fileopenbox(filetypes=['*.docx'])
# if targetFile=="":
#     exit()
#print(targetFile)
#targetFile=open(str(targetFile),"r",errors='ignore')
#text=targetFile.read()
fullText =getText(targetFile) #docx2txt.process(str(targetFile))


#print(text)
httpClient = None


doc=docx.Document()

saveFile=eg.filesavebox(filetypes=['*.docx'],default="newTranslated")
if saveFile==None:
    sys.exit()

for para in fullText:
    
    print(para)
    myurl = '/api/trans/vip/translate'
    md5 = hashlib.md5()
    fromLang = 'en'
    toLang = 'zh'
    q=para#'hello baidu'
    salt = random.randint(32768, 65536)
    sign = appid+q+str(salt)+secretKey
    md5.update(sign.encode("utf8"))
    sign = md5.hexdigest()
    # print(u'\u6211'.encode('utf-8').decode('utf8'))
    myurl = myurl+'?appid='+appid+'&q='+q+'&from='+fromLang+'&to='+toLang+'&salt='+str(salt)+'&sign='+sign
    # print (myurl)
    myurl=' http://api.fanyi.baidu.com'+myurl
    # print(myurl)
    # print(requests.get(myurl).json())
    response=requests.get(myurl).json().get('trans_result')
    if response!=None:
    
        for para in response:
            doc.add_paragraph(para["src"])
            doc.add_paragraph(para["dst"]+"\n")
    time.sleep(1)
    

doc.save(str(saveFile)+'.docx')
